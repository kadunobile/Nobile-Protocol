import re
import time
import logging
from functools import lru_cache
from typing import Optional, List

import streamlit as st
import PyPDF2
import streamlit.components.v1 as components
from openai import OpenAI, APITimeoutError, RateLimitError

from core.data import CIDADES_BRASIL

# Configurar logger para este módulo
logger = logging.getLogger(__name__)

def corrigir_formatacao(texto: Optional[str]) -> Optional[str]:
    """
    Corrige a formatação de texto removendo prefixos de moeda e normalizando espaços.
    
    Args:
        texto: Texto a ser formatado
        
    Returns:
        Texto formatado ou None se entrada for None
        
    Examples:
        >>> corrigir_formatacao("R$ 5.000")
        "5.000"
        >>> corrigir_formatacao(None)
        None
    """
    if not texto:
        return texto

    # Remove qualquer prefixo de moeda (R$ ou R)
    texto = re.sub(r'\bR\$\s*', '', texto)
    texto = re.sub(r'\bR\s+', '', texto)

    # Mantém quebras de linha (não usar \s aqui)
    texto = re.sub(r'[ \t]{2,}', ' ', texto)

    # Normaliza quebras do Windows
    texto = texto.replace('\r\n', '\n')

    return texto

@st.cache_data(ttl=3600)
def extrair_texto_pdf(arquivo) -> Optional[str]:
    """
    Extrai texto de um arquivo PDF.
    
    Utiliza cache de 1 hora para evitar reprocessamento do mesmo arquivo.
    
    Args:
        arquivo: Objeto de arquivo PDF (UploadedFile do Streamlit)
        
    Returns:
        Texto extraído do PDF ou None em caso de erro
        
    Raises:
        Nenhuma exceção é propagada; erros são logados e exibidos ao usuário
    """
    logger.info("Iniciando extração de texto de PDF")
    
    try:
        pdf_reader = PyPDF2.PdfReader(arquivo)
        num_pages = len(pdf_reader.pages)
        logger.debug(f"PDF possui {num_pages} páginas")
        
        texto = "".join([p.extract_text() for p in pdf_reader.pages])
        
        # Validar se o arquivo tem conteúdo
        if not texto or not texto.strip():
            logger.warning("PDF não contém texto extraível")
            st.error("O PDF não contém texto extraível. Certifique-se de que não é uma imagem escaneada.")
            return None
        
        logger.info(f"Texto extraído com sucesso ({len(texto)} caracteres)")
        return texto
        
    except PyPDF2.errors.PdfReadError as e:
        logger.error(f"Erro ao ler PDF (arquivo corrompido ou inválido): {e}")
        st.error(f"Erro ao ler PDF: arquivo pode estar corrompido ou protegido por senha")
        return None
    except Exception as e:
        logger.error(f"Erro inesperado ao extrair texto do PDF: {e}", exc_info=True)
        st.error(f"Erro ao ler PDF: {e}")
        return None

def extrair_texto_docx(arquivo):
    """
    Extrai texto de arquivo Word (.docx).
    
    Args:
        arquivo: Arquivo uploaded do Streamlit
    
    Returns:
        str: Texto extraído ou None em caso de erro
    """
    logger = logging.getLogger(__name__)
    logger.info("Iniciando extração de texto de DOCX")
    
    try:
        import docx
        doc = docx.Document(arquivo)
        
        # Extrai todos os parágrafos
        texto_completo = []
        for paragrafo in doc.paragraphs:
            if paragrafo.text.strip():
                texto_completo.append(paragrafo.text)
        
        # Extrai tabelas (comum em CVs)
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    if celula.text.strip():
                        texto_completo.append(celula.text)
        
        texto = "\n".join(texto_completo)
        
        if not texto.strip():
            st.error("⚠️ Arquivo Word vazio ou sem texto extraível")
            logger.warning("DOCX extraído está vazio")
            return None
        
        logger.info(f"DOCX extraído com sucesso: {len(texto)} caracteres")
        return texto
        
    except Exception as e:
        st.error(f"❌ Erro ao ler arquivo Word: {e}")
        logger.error(f"Erro ao extrair DOCX: {e}", exc_info=True)
        return None


def extrair_texto_txt(arquivo):
    """
    Extrai texto de arquivo TXT.
    
    Args:
        arquivo: Arquivo uploaded do Streamlit
    
    Returns:
        str: Texto extraído ou None em caso de erro
    """
    logger = logging.getLogger(__name__)
    logger.info("Iniciando extração de texto de TXT")
    
    try:
        # Tenta diferentes encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                arquivo.seek(0)  # Reset file pointer
                texto = arquivo.read().decode(encoding)
                
                if not texto.strip():
                    st.error("⚠️ Arquivo TXT vazio")
                    logger.warning("TXT extraído está vazio")
                    return None
                
                logger.info(f"TXT extraído com sucesso ({encoding}): {len(texto)} caracteres")
                return texto
                
            except (UnicodeDecodeError, AttributeError):
                continue
        
        # Se nenhum encoding funcionou
        st.error("❌ Não foi possível decodificar o arquivo TXT")
        logger.error("Falha em todos os encodings para TXT")
        return None
        
    except Exception as e:
        st.error(f"❌ Erro ao ler arquivo TXT: {e}")
        logger.error(f"Erro ao extrair TXT: {e}", exc_info=True)
        return None


def extrair_texto_universal(arquivo, tipo_arquivo):
    """
    Função universal que detecta o tipo e extrai o texto.
    
    Args:
        arquivo: Arquivo uploaded do Streamlit
        tipo_arquivo: Extensão do arquivo (ex: 'pdf', 'docx', 'txt')
    
    Returns:
        str: Texto extraído ou None em caso de erro
    """
    tipo_arquivo = tipo_arquivo.lower()
    
    if tipo_arquivo == 'pdf':
        return extrair_texto_pdf(arquivo)
    elif tipo_arquivo in ['docx', 'doc']:
        return extrair_texto_docx(arquivo)
    elif tipo_arquivo == 'txt':
        return extrair_texto_txt(arquivo)
    else:
        st.error(f"❌ Formato não suportado: {tipo_arquivo}")
        return None

def inicializar_cliente_openai(key: str) -> Optional[OpenAI]:
    """
    Inicializa o cliente OpenAI com validação de API key.
    
    Args:
        key: Chave da API OpenAI
        
    Returns:
        Cliente OpenAI inicializado ou None em caso de erro
        
    Raises:
        Nenhuma exceção é propagada; erros são logados e exibidos ao usuário
    """
    logger.info("Inicializando cliente OpenAI")
    
    # Validar formato da API key
    if not key or not isinstance(key, str):
        logger.error("API key não fornecida ou inválida")
        st.error("API key não fornecida")
        return None
    
    if not key.startswith("sk-"):
        logger.warning("API key não começa com 'sk-' (formato esperado)")
        st.warning("API key parece estar em formato inválido (deve começar com 'sk-')")
    
    try:
        client = OpenAI(api_key=key)
        
        # Fazer teste básico de conexão (list models é rápido e barato)
        try:
            # Apenas criar o cliente é suficiente - validação real será feita no primeiro uso
            logger.info("Cliente OpenAI inicializado com sucesso")
            return client
        except Exception as e:
            logger.error(f"Erro ao testar conexão com OpenAI: {e}")
            st.error(f"Erro ao validar conexão com OpenAI: {e}")
            return None
            
    except Exception as e:
        logger.error(f"Erro ao inicializar cliente OpenAI: {e}", exc_info=True)
        st.error(f"Erro ao conectar com OpenAI: Verifique sua API key")
        return None

def chamar_gpt(client: OpenAI, msgs: list, max_retries: int = 3, timeout: int = 30) -> Optional[str]:
    """
    Chama a API do GPT com retry automático e tratamento robusto de erros.
    
    Implementa backoff exponencial para timeouts e trata especificamente
    erros de rate limit e timeout.
    
    Args:
        client: Cliente OpenAI inicializado
        msgs: Lista de mensagens no formato esperado pela API
        max_retries: Número máximo de tentativas (padrão: 3)
        timeout: Timeout em segundos por requisição (padrão: 30)
        
    Returns:
        Resposta do GPT formatada ou None em caso de erro
        
    Raises:
        Nenhuma exceção é propagada; erros são logados e exibidos ao usuário
        
    Examples:
        >>> client = OpenAI(api_key="sk-...")
        >>> msgs = [{"role": "user", "content": "Hello"}]
        >>> resposta = chamar_gpt(client, msgs)
    """
    logger.info(f"Chamando GPT com {len(msgs)} mensagens")
    
    for tentativa in range(1, max_retries + 1):
        try:
            logger.debug(f"Tentativa {tentativa}/{max_retries}")
            
            response = client.chat.completions.create(
                model="gpt-4o",
                messages=msgs,
                temperature=0.7,
                max_tokens=4000,
                timeout=timeout
            )
            
            texto_raw = response.choices[0].message.content
            logger.info(f"Resposta recebida com sucesso ({len(texto_raw)} caracteres)")
            return corrigir_formatacao(texto_raw)
            
        except APITimeoutError as e:
            logger.warning(f"Timeout na tentativa {tentativa}/{max_retries}: {e}")
            
            if tentativa < max_retries:
                # Backoff exponencial: 2^tentativa segundos
                tempo_espera = 2 ** tentativa
                logger.info(f"Aguardando {tempo_espera}s antes de tentar novamente...")
                time.sleep(tempo_espera)
            else:
                logger.error("Todas as tentativas falharam por timeout")
                st.error("Erro: Timeout ao chamar GPT. Tente novamente mais tarde.")
                return None
                
        except RateLimitError as e:
            logger.error(f"Rate limit atingido: {e}")
            st.error("Erro: Limite de requisições atingido. Por favor, aguarde alguns minutos e tente novamente.")
            return None
            
        except Exception as e:
            logger.error(f"Erro inesperado na tentativa {tentativa}/{max_retries}: {e}", exc_info=True)
            
            if tentativa < max_retries:
                tempo_espera = 2 ** tentativa
                logger.info(f"Aguardando {tempo_espera}s antes de tentar novamente...")
                time.sleep(tempo_espera)
            else:
                logger.error("Todas as tentativas falharam")
                st.error(f"Erro ao chamar GPT: {e}")
                return None
    
    return None

def scroll_topo():
    components.html("""
        <script>
            const scrollToTop = () => {
                const main = window.parent.document.querySelector('section.main');
                if (main) main.scrollTop = 0;
                window.parent.scrollTo(0, 0);
                document.documentElement.scrollTop = 0;
                document.body.scrollTop = 0;
            };
            scrollToTop();
            setTimeout(scrollToTop, 50);
            setTimeout(scrollToTop, 150);
            setTimeout(scrollToTop, 400);
        </script>
    """, height=0)

@lru_cache(maxsize=100)
def filtrar_cidades(texto: Optional[str]) -> tuple:
    """
    Filtra cidades brasileiras com base em texto de busca.
    
    Utiliza cache LRU para otimizar buscas repetidas.
    
    Args:
        texto: Texto para filtrar cidades (case-insensitive)
        
    Returns:
        Tupla de cidades que correspondem ao filtro (limitado a 10 se texto vazio)
        
    Examples:
        >>> filtrar_cidades("São")
        ('São Paulo/SP', 'São José/SC', ...)
        >>> filtrar_cidades(None)
        ('São Paulo/SP', 'Rio de Janeiro/RJ', ...)  # primeiras 10
    """
    if not texto:
        return tuple(CIDADES_BRASIL[:10])
    resultado = [c for c in CIDADES_BRASIL if texto.lower() in c.lower()]
    return tuple(resultado)

def forcar_topo():
    components.html("""
        <script>
            const goTop = () => {
                const main = window.parent.document.querySelector('section.main');
                if (main) main.scrollTop = 0;
                window.parent.scrollTo(0, 0);
                document.documentElement.scrollTop = 0;
                document.body.scrollTop = 0;
            };

            const blurInputs = () => {
                const inputs = window.parent.document.querySelectorAll('textarea, input');
                inputs.forEach(i => i.blur());
            };

            goTop();
            blurInputs();

            setTimeout(() => { goTop(); blurInputs(); }, 50);
            setTimeout(() => { goTop(); blurInputs(); }, 150);
            setTimeout(() => { goTop(); blurInputs(); }, 400);
        </script>
    """, height=0)