"""
Exportador de CV em formato DOCX.

Este módulo gera currículos formatados profissionalmente em Microsoft Word.
"""

import io
import logging
from typing import Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configurar logger
logger = logging.getLogger(__name__)


def gerar_cv_docx(dados_cv: dict) -> io.BytesIO:
    """
    Gera CV formatado em DOCX.
    
    Cria um documento Word profissional com seções estruturadas,
    formatação consistente e layout otimizado para ATS.
    
    Args:
        dados_cv: Dict com seções do CV:
            - nome: str
            - email: str (opcional)
            - telefone: str (opcional)
            - linkedin: str (opcional)
            - resumo: str (opcional)
            - experiencias: List[Dict] (opcional)
                - cargo: str
                - empresa: str
                - periodo: str
                - descricao: str (opcional)
                - realizacoes: List[str] (opcional)
            - educacao: List[Dict] (opcional)
                - curso: str
                - instituicao: str
                - periodo: str
            - habilidades: Dict[str, List[str]] (opcional)
    
    Returns:
        BytesIO com o documento Word
        
    Raises:
        Exception: Se houver erro na geração do documento
        
    Examples:
        >>> dados = {'nome': 'João Silva', 'email': 'joao@email.com'}
        >>> buffer = gerar_cv_docx(dados)
        >>> # buffer pode ser salvo ou enviado como download
    """
    logger.info("Iniciando geração de CV em DOCX")
    
    try:
        doc = Document()
        
        # ===== CONFIGURAR MARGENS =====
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # ===== CABEÇALHO - NOME =====
        nome = doc.add_heading(dados_cv.get('nome', 'Nome Completo'), level=1)
        nome.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nome.runs[0].font.color.rgb = RGBColor(0, 51, 102)  # Azul escuro profissional
        nome.runs[0].font.size = Pt(24)
        
        # ===== INFORMAÇÕES DE CONTATO =====
        contato_parts = []
        if dados_cv.get('email'):
            contato_parts.append(dados_cv['email'])
        if dados_cv.get('telefone'):
            contato_parts.append(dados_cv['telefone'])
        if dados_cv.get('linkedin'):
            contato_parts.append(dados_cv['linkedin'])
        
        if contato_parts:
            contato = doc.add_paragraph()
            contato.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contato_text = ' | '.join(contato_parts)
            run = contato.add_run(contato_text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(64, 64, 64)  # Cinza escuro
        
        # Espaço após cabeçalho
        doc.add_paragraph()
        
        # ===== RESUMO PROFISSIONAL =====
        if dados_cv.get('resumo'):
            heading = doc.add_heading('RESUMO PROFISSIONAL', level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            
            resumo_p = doc.add_paragraph(dados_cv['resumo'])
            resumo_p.paragraph_format.space_after = Pt(12)
        
        # ===== EXPERIÊNCIAS PROFISSIONAIS =====
        if dados_cv.get('experiencias'):
            heading = doc.add_heading('EXPERIÊNCIA PROFISSIONAL', level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            
            for exp in dados_cv['experiencias']:
                # Cargo e Empresa
                p = doc.add_paragraph()
                cargo_run = p.add_run(f"{exp.get('cargo', 'Cargo')} | {exp.get('empresa', 'Empresa')}")
                cargo_run.bold = True
                cargo_run.font.size = Pt(11)
                
                # Período
                periodo = doc.add_paragraph()
                periodo_run = periodo.add_run(exp.get('periodo', ''))
                periodo_run.italic = True
                periodo_run.font.size = Pt(10)
                periodo_run.font.color.rgb = RGBColor(96, 96, 96)
                periodo.paragraph_format.space_after = Pt(6)
                
                # Descrição
                if exp.get('descricao'):
                    desc_p = doc.add_paragraph(exp['descricao'])
                    desc_p.paragraph_format.space_after = Pt(6)
                
                # Realizações com bullets
                if exp.get('realizacoes'):
                    for realizacao in exp['realizacoes']:
                        p = doc.add_paragraph(realizacao, style='List Bullet')
                        p.paragraph_format.space_after = Pt(3)
                        p.paragraph_format.left_indent = Inches(0.25)
                
                # Espaço entre experiências
                doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        # ===== FORMAÇÃO ACADÊMICA =====
        if dados_cv.get('educacao'):
            heading = doc.add_heading('FORMAÇÃO ACADÊMICA', level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            
            for edu in dados_cv['educacao']:
                # Curso e Instituição
                p = doc.add_paragraph()
                curso_run = p.add_run(f"{edu.get('curso', 'Curso')} - {edu.get('instituicao', 'Instituição')}")
                curso_run.bold = True
                curso_run.font.size = Pt(11)
                
                # Período
                periodo = doc.add_paragraph(edu.get('periodo', ''))
                periodo.runs[0].font.size = Pt(10)
                periodo.runs[0].font.color.rgb = RGBColor(96, 96, 96)
                periodo.paragraph_format.space_after = Pt(12)
        
        # ===== HABILIDADES =====
        if dados_cv.get('habilidades'):
            heading = doc.add_heading('HABILIDADES', level=2)
            heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
            
            # Se habilidades é um dict (categorizado)
            if isinstance(dados_cv['habilidades'], dict):
                for categoria, skills in dados_cv['habilidades'].items():
                    p = doc.add_paragraph()
                    categoria_run = p.add_run(f"{categoria}: ")
                    categoria_run.bold = True
                    categoria_run.font.size = Pt(10)
                    
                    skills_text = ', '.join(skills) if isinstance(skills, list) else str(skills)
                    skills_run = p.add_run(skills_text)
                    skills_run.font.size = Pt(10)
                    p.paragraph_format.space_after = Pt(6)
            
            # Se habilidades é uma lista simples
            elif isinstance(dados_cv['habilidades'], list):
                p = doc.add_paragraph(', '.join(dados_cv['habilidades']))
                p.paragraph_format.space_after = Pt(6)
        
        # ===== SALVAR EM BYTESIO =====
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info("CV DOCX gerado com sucesso")
        return buffer
        
    except Exception as e:
        logger.error(f"Erro ao gerar CV DOCX: {e}", exc_info=True)
        raise
